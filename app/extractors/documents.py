import io
import pdfplumber
import docx
import pytesseract
from app.extractors.base import BaseExtractor
from app.schemas import ExtractedUnit, Location
from app.utils import preprocess_image_for_ocr

class PDFExtractor(BaseExtractor):
    def extract(self):
        units = []
        with pdfplumber.open(io.BytesIO(self.file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                
                # FALLBACK: If text is empty/garbage, assume scanned PDF
                if not text or len(text.strip()) < 10:
                    # Render page as image for OCR
                    im = page.to_image(resolution=300).original
                    # Convert PIL image to bytes to reuse our util
                    buf = io.BytesIO()
                    im.save(buf, format='PNG')
                    processed_im = preprocess_image_for_ocr(buf.getvalue())
                    text = pytesseract.image_to_string(processed_im)

                if text and text.strip():
                    units.append(ExtractedUnit(
                        text=text.strip(),
                        source=f"page_{i+1}",
                        location=Location(type="page", number=i+1)
                    ))
        return units

# In app/extractors/documents.py

class WordExtractor(BaseExtractor):
    def extract(self):
        doc = docx.Document(io.BytesIO(self.file_bytes))
        units = []
        
        # 1. Extract Standard Paragraphs
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                units.append(ExtractedUnit(
                    text=para.text.strip(),
                    source="paragraph",
                    location=Location(type="row", number=i+1)
                ))

        # 2. Extract Text from Tables (The Missing Piece)
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                # specific logic: join cell text so it reads like a line
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                
                if row_text:
                    units.append(ExtractedUnit(
                        text=row_text,
                        source=f"table_{t_idx+1}",
                        location=Location(type="row", number=r_idx+1)
                    ))
                    
        return units