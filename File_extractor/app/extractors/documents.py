import io
import pdfplumber
import docx
import pytesseract
from app.extractors.base import BaseExtractor
from app.schemas import ExtractedUnit, Location
from app.utils import preprocess_image_for_ocr

class PDFExtractor(BaseExtractor):
    def extract(self):
        units = []
        with pdfplumber.open(io.BytesIO(self.file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                # 1. Try standard text extraction
                text = page.extract_text()
                
                # Check if page is readable text (Native PDF)
                is_native_text = text and len(text.strip()) >= 10

                if is_native_text:
                    # A. Add the text found
                    units.append(ExtractedUnit(
                        text=text.strip(),
                        source=f"page_{i+1}",
                        location=Location(type="page", number=i+1)
                    ))

                    # B. Look for embedded images "in between" the text
                    # (e.g., charts, photos in a digital PDF)
                    for img_idx, img in enumerate(page.images):
                        try:
                            # Filter out tiny artifacts (lines/icons < 50px)
                            x0, top, x1, bottom = img['x0'], img['top'], img['x1'], img['bottom']
                            if (x1 - x0) < 50 or (bottom - top) < 50:
                                continue

                            # Crop and OCR
                            cropped = page.crop((x0, top, x1, bottom))
                            # Convert to high-res image for OCR
                            im_obj = cropped.to_image(resolution=300).original
                            
                            # Convert to bytes for our preprocessor
                            buf = io.BytesIO()
                            im_obj.save(buf, format='PNG')
                            
                            processed_im = preprocess_image_for_ocr(buf.getvalue())
                            img_text = pytesseract.image_to_string(processed_im).strip()

                            if img_text:
                                units.append(ExtractedUnit(
                                    text=f"[Image Extraction]: {img_text}",
                                    source=f"page_{i+1}_img_{img_idx+1}",
                                    location=Location(type="pixel_box", number=i+1)
                                ))
                        except Exception as e:
                            # If an image fails, skip it and continue
                            print(f"PDF Image Extract Error page {i+1}: {e}")

                else:
                    # FALLBACK: Page is empty or looks like a scan
                    # Render entire page as image for OCR
                    im = page.to_image(resolution=300).original
                    buf = io.BytesIO()
                    im.save(buf, format='PNG')
                    processed_im = preprocess_image_for_ocr(buf.getvalue())
                    text = pytesseract.image_to_string(processed_im)
                    
                    if text and text.strip():
                        units.append(ExtractedUnit(
                            text=text.strip(),
                            source=f"page_{i+1}_full_ocr",
                            location=Location(type="page", number=i+1)
                        ))

        return units


class WordExtractor(BaseExtractor):
    def extract(self):
        doc = docx.Document(io.BytesIO(self.file_bytes))
        units = []
        
        # 1. Extract Paragraphs & Inline Images
        for i, para in enumerate(doc.paragraphs):
            # A. Standard Text
            if para.text.strip():
                units.append(ExtractedUnit(
                    text=para.text.strip(),
                    source="paragraph",
                    location=Location(type="row", number=i+1)
                ))
            
            # B. Check for Embedded Images in this paragraph
            # We iterate through 'runs' to find XML drawing elements
            for run in para.runs:
                if 'drawing' in run.element.xml:
                    try:
                        # Iterate XML children to find the image reference (blip)
                        for elem in run.element.iter():
                            if elem.tag.endswith('blip'):
                                # Get the Relationship ID (r:embed)
                                embed_attr = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                
                                if embed_attr and embed_attr in doc.part.rels:
                                    # Get the image data
                                    image_part = doc.part.rels[embed_attr].target_part
                                    image_bytes = image_part.blob
                                    
                                    # Run OCR
                                    processed = preprocess_image_for_ocr(image_bytes)
                                    img_text = pytesseract.image_to_string(processed).strip()
                                    
                                    if img_text:
                                        units.append(ExtractedUnit(
                                            text=f"[Image Extraction]: {img_text}",
                                            source=f"inline_image_para_{i+1}",
                                            location=Location(type="row", number=i+1)
                                        ))
                    except Exception as e:
                        print(f"Word Image Extract Error: {e}")

        # 2. Extract Text from Tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                
                if row_text:
                    units.append(ExtractedUnit(
                        text=row_text,
                        source=f"table_{t_idx+1}",
                        location=Location(type="row", number=r_idx+1)
                    ))
                    
        return units